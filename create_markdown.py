from docx import Document
import pdfplumber
import argparse
import re
import os
import sys
import io
import json
import time
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from openai import OpenAI
from dotenv import load_dotenv


# =========================
# 1. 读取 .env 配置
# =========================
load_dotenv()

GLM_API_KEY = os.getenv("API_KEY")
GLM_BASE_URL = os.getenv("API_BASE_URL")
GLM_MODEL = os.getenv("API_MODEL") or "glm-latest"

client = OpenAI(
    api_key=GLM_API_KEY,
    base_url=GLM_BASE_URL
)

DATA_DIR = "data"
MARKDOWN_DIR = os.path.join(DATA_DIR, "markdown")
os.makedirs(MARKDOWN_DIR, exist_ok=True)

# 日志配置：同时输出到文件（UTF-8，不受控制台编码影响）和控制台。
# 控制台 handler 用一个能容错编码的包装，避免在 GBK 控制台（部分 Windows Server）
# 打印中文时抛 UnicodeEncodeError（表现为刷屏 "Logging error ... in emit"）。
_log_formatter = logging.Formatter(
    fmt="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)

# 文件 handler：始终 UTF-8，保证完整日志落盘可回溯
_file_handler = logging.FileHandler(
    os.path.join(DATA_DIR, "run_markdown.log"), mode="a", encoding="utf-8"
)
_file_handler.setFormatter(_log_formatter)

# 控制台 handler：强制以 UTF-8 输出。优先 reconfigure，失败则用 TextIOWrapper
# 直接包装底层 buffer，彻底绕开 Python 判定的 GBK 编码（部分 Windows Server 上
# reconfigure 不足以生效，会刷屏 "Logging error ... in emit"）。
_console_stream = sys.stdout
try:
    if hasattr(_console_stream, "reconfigure"):
        _console_stream.reconfigure(encoding="utf-8", errors="replace")
    elif hasattr(_console_stream, "buffer"):
        _console_stream = io.TextIOWrapper(
            _console_stream.buffer, encoding="utf-8", errors="replace", line_buffering=True
        )
except Exception:
    # 兜底：即便包装失败，也让控制台 handler 用容错编码，不影响文件日志
    if hasattr(sys.stdout, "buffer"):
        try:
            _console_stream = io.TextIOWrapper(
                sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True
            )
        except Exception:
            _console_stream = sys.stdout
_console_handler = logging.StreamHandler(_console_stream)
_console_handler.setFormatter(_log_formatter)

logging.basicConfig(level=logging.INFO, handlers=[_file_handler, _console_handler])
logger = logging.getLogger("create_markdown")

GENERATION_PROMPT =''' 
【角色设定】
你是失效分析知识工程师，专长于核电厂设备失效报告的结构化提取与知识库入库预处理。

【任务目标】
将输入的失效分析技术报告内容，提取为Markdown格式。

【强制提取字段】（必须完整填充，无信息则标注"未提及"）
1. **报告标识**：原始报告完整名称（剔除日期信息）
2. **失效部件**：
   - 部件名称、材料牌号/等级（如35CrMo、8.8级、Inconel X-750）
   - 规格参数（尺寸、硬度、强度等级）
   - 在设备中的具体位置/功能
3. **场景信息**：
   - 电厂名称（如红沿河核电厂）
   - 系统代码及全称（如CEX-凝结水抽取系统）
   - 设备位号（如H5SIR221PO）
   - 发生时间（精确到年月，含服役时长）
4. **失效机理**（200字内）：
   - 失效性质（氢脆/疲劳/脆性/腐蚀/过载等）
   - 关键检验数据（ 动态提取，禁止限定固定章节：宏观检查、成分分析、金相检验、力学性能、微观特征、能谱分析、腐蚀产物成分等）
   
5. **失效结论**： 整段保留，不要修改任何字
6. **建议措施**： 整段保留，不要修改任何字
7. **经验标签**：
   - 失效模式标签（如氢脆断裂、旋转弯曲疲劳、点蚀）
   - 材料类别（合金钢/碳钢/铸铁/镍基合金/不锈钢）
   - 系统分类
   - 设备类型（泵/阀门/齿轮箱/弹簧等）

【格式规范】
- 输出为Markdown代码块，使用以下结构：
```markdown
# [报告主标题]

## 基本信息
- **报告名称**: 
- **失效部件**: [名称]，[材料]，[规格]，[位置]
- **所属电厂**: 
- **所属系统**: [代码]-[全称]
- **发生时间**: [发现时间]，服役[时长]

## 失效分析结果

### 关键检验数据
- **宏观检查**: 
- **成分分析**: 
- **金相组织**: 
- **力学性能**: [硬度/强度具体数值及与标准对比]
- **微观特征**: 
- **数据标准化要求**：
     * 所有数值必须标注**单位**（如**15.7%**、**255HV10**、**28J**）
     * 必须包含**标准符合性判定**（符合/超标/不合格/未提及）
     * 裂纹/断口必须描述**启裂位置**和**扩展方向**


### 失效机理
[机理描述，突出关键数据]

## 失效结论
[结论]

## 后续建议
- [建议]

## 经验标签
`失效模式:XXX` `材料:XXX` `系统:XXX` `设备:XXX`

【动态提取规则】
1. **检测项目识别**：通读全文，识别所有检测章节（如"晶间腐蚀性能检验"、"氢含量检测"等），禁止遗漏
2. **数据优先级**：
   - 优先提取**与失效直接相关**的检测数据（如断口分析、腐蚀产物成分）
   - 次要提取**材料基础性能**（如常规拉伸、硬度）
3. **缺失处理**：若报告中某检测项目未开展，标注"未检测"而非省略
4. **单位统一**：确保所有数值带单位，百分比保留1位小数，硬度标注标尺（HRC/HV/HBW）

【长度控制规则】

除了失效结论和建议措施，单个字段描述不超过300字，避免向量化截断
数值类信息优先使用加粗突出（如硬度**520HV**）
删除原文中的示意图描述、寒暄语句、标准全文引用，仅保留结论性数据

【质量检查】

输出前自检：是否包含以下关键数值？
[ ] 材料牌号/等级
[ ] 具体硬度/强度数值（如**520HV**、**1700MPa**）
[ ] 服役时间（如8个月、25个月）
[ ] 失效性质判定关键词（氢脆/疲劳/脆性/腐蚀）
[ ] 标准符合性判定（符合/超标/不合格）
'''


def extract_docx_content(docx_path: str) -> str:
    """读取 Word 正文和表格内容，不做清洗。"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise ValueError(
            f"无法作为 .docx 打开（可能是老式 .doc 改名、损坏或空文件）：{e}"
        )

    content_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            content_parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            content_parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(content_parts)


def extract_pdf_content(pdf_path: str) -> str:
    """读取 PDF 全部页面文本，不做清洗。"""
    try:
        pdf = pdfplumber.open(pdf_path)
    except Exception as e:
        raise ValueError(f"无法作为 PDF 打开（文件可能损坏、加密或为空）：{e}")

    with pdf:
        return "\n\n".join(page.extract_text() or "" for page in pdf.pages)


def call_glm_generate(article_text: str, max_retries: int = 3) -> str:
    """将整篇文章一次性输入模型并返回 Markdown。"""
    user_content = f"""
{GENERATION_PROMPT}

以下是未经清洗、未经分块的原始全文：

{article_text}
"""
    messages = [{"role": "user", "content": user_content}]

    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=GLM_MODEL,
                messages=messages,
                temperature=0,
                timeout=1800.0,
            )
            content = resp.choices[0].message.content
            if content is None or not content.strip():
                raise ValueError("模型返回空内容")
            return content.strip()
        except Exception as e:
            last_err = e
            if attempt < max_retries:
                sleep_s = min(2 ** attempt, 30)
                logger.warning(
                    "模型调用失败（第 %d/%d 次）：%s，%d 秒后重试",
                    attempt,
                    max_retries,
                    e,
                    sleep_s,
                )
                time.sleep(sleep_s)

    raise last_err


def strip_markdown_fence(content: str) -> str:
    """移除模型可能添加的最外层 Markdown 代码围栏。"""
    lines = content.strip().splitlines()
    if len(lines) >= 2 and lines[0].strip().lower() in ("```", "```markdown", "```md"):
        if lines[-1].strip() == "```":
            return "\n".join(lines[1:-1]).strip()
    return content.strip()


def process_one_file(file_path: str, output_path: str, max_retries: int) -> dict:
    """读取一篇报告，单次调用模型，并保存为一个 Markdown 文件。"""
    filename = os.path.basename(file_path)
    if filename.lower().endswith(".pdf"):
        text = extract_pdf_content(file_path)
    else:
        text = extract_docx_content(file_path)

    if not text.strip():
        return {"filename": filename, "empty": True, "chars": 0}

    logger.info("[%s] 全文读取完成（%d 字），开始单次模型调用", filename, len(text))
    markdown = strip_markdown_fence(
        call_glm_generate(text, max_retries=max_retries)
    )

    with open(output_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(markdown)
        f.write("\n")

    return {"filename": filename, "empty": False, "chars": len(text)}


def load_generated_files(log_path: str) -> set:
    """读取已经成功生成 Markdown 的源文件名。"""
    if not os.path.exists(log_path):
        return set()

    with open(log_path, "r", encoding="utf-8") as f:
        return {line.strip() for line in f if line.strip()}


def mark_generated(log_path: str, filename: str):
    """将成功生成的源文件名追加到日志。"""
    with open(log_path, "a", encoding="utf-8", newline="\n") as f:
        f.write(filename + "\n")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input_folder", default="papers_failure_analysis2")
    parser.add_argument("--output_folder", default=MARKDOWN_DIR)
    parser.add_argument(
        "--error_file", default=os.path.join(DATA_DIR, "error_log_markdown.txt")
    )
    parser.add_argument("--max_workers", type=int, default=4, help="文件间并发数")
    parser.add_argument("--max_retries", type=int, default=3, help="单次模型调用最大尝试次数")
    args = parser.parse_args()

    if not GLM_API_KEY or not GLM_BASE_URL:
        logger.error("API_KEY / API_BASE_URL 未配置，请检查 .env 文件")
        return
    if not os.path.isdir(args.input_folder):
        logger.error("输入目录不存在：%s", args.input_folder)
        return

    os.makedirs(args.output_folder, exist_ok=True)
    markdown_log = os.path.join(args.output_folder, "markdown_log.txt")
    generated_files = load_generated_files(markdown_log)
    filenames = sorted(
        filename
        for filename in os.listdir(args.input_folder)
        if filename.lower().endswith((".docx", ".pdf"))
        and not filename.startswith("~$")
    )

    output_names = [os.path.splitext(filename)[0] + ".md" for filename in filenames]
    duplicates = sorted({name for name in output_names if output_names.count(name) > 1})
    if duplicates:
        logger.error("存在同名 Word/PDF，无法安全生成 Markdown：%s", ", ".join(duplicates))
        return

    jobs = []
    for filename, output_name in zip(filenames, output_names):
        output_path = os.path.join(args.output_folder, output_name)
        if filename in generated_files:
            logger.info("[%s] 已在生成日志中，跳过", filename)
            continue
        if os.path.exists(output_path):
            mark_generated(markdown_log, filename)
            generated_files.add(filename)
            logger.info("[%s] Markdown 已存在，已补记生成日志并跳过", filename)
            continue
        jobs.append((filename, output_path))

    if not jobs:
        logger.info("没有需要处理的文件（共 %d 个）", len(filenames))
        return

    logger.info("待处理 %d 个文件，并发数 %d", len(jobs), args.max_workers)
    succeeded = 0
    with ThreadPoolExecutor(max_workers=args.max_workers) as executor:
        future_to_job = {
            executor.submit(
                process_one_file,
                os.path.join(args.input_folder, filename),
                output_path,
                args.max_retries,
            ): (filename, output_path)
            for filename, output_path in jobs
        }

        for future in as_completed(future_to_job):
            filename, output_path = future_to_job[future]
            try:
                result = future.result()
                if result["empty"]:
                    logger.warning("[%s] 提取文本为空，跳过", filename)
                    continue
                mark_generated(markdown_log, filename)
                generated_files.add(filename)
                succeeded += 1
                logger.info("[%s] 完成 -> %s", filename, output_path)
            except Exception as e:
                logger.error("[%s] 处理失败：%s", filename, e)
                with open(args.error_file, "a", encoding="utf-8") as f:
                    f.write(f"{filename}\t{str(e)}\n")

    logger.info("全部结束，本次成功生成 %d 个 Markdown 文件", succeeded)


if __name__ == "__main__":
    main()
