from docx import Document
import argparse
import re
import os
import sys
import io
import json
import time
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from openai import OpenAI
from dotenv import load_dotenv


# =========================
# 1. 读取 .env 配置
# =========================
load_dotenv()

GLM_API_KEY = os.getenv("API_KEY")
GLM_BASE_URL = os.getenv("API_BASE_URL")
GLM_MODEL = os.getenv("API_MODEL") or "glm-latest"

client = OpenAI(
    api_key=GLM_API_KEY,
    base_url=GLM_BASE_URL
)

MIN_CHUNK_CHARS = 6000 
MAX_CHUNK_CHARS = 12000

# 日志配置：同时输出到文件（UTF-8，不受控制台编码影响）和控制台。
# 控制台 handler 用一个能容错编码的包装，避免在 GBK 控制台（部分 Windows Server）
# 打印中文时抛 UnicodeEncodeError（表现为刷屏 "Logging error ... in emit"）。
_log_formatter = logging.Formatter(
    fmt="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)

# 文件 handler：始终 UTF-8，保证完整日志落盘可回溯
_file_handler = logging.FileHandler("data/run_word.log", mode="a", encoding="utf-8")
_file_handler.setFormatter(_log_formatter)

# 控制台 handler：强制以 UTF-8 输出。优先 reconfigure，失败则用 TextIOWrapper
# 直接包装底层 buffer，彻底绕开 Python 判定的 GBK 编码（部分 Windows Server 上
# reconfigure 不足以生效，会刷屏 "Logging error ... in emit"）。
_console_stream = sys.stdout
try:
    if hasattr(_console_stream, "reconfigure"):
        _console_stream.reconfigure(encoding="utf-8", errors="replace")
    elif hasattr(_console_stream, "buffer"):
        _console_stream = io.TextIOWrapper(
            _console_stream.buffer, encoding="utf-8", errors="replace", line_buffering=True
        )
except Exception:
    # 兜底：即便包装失败，也让控制台 handler 用容错编码，不影响文件日志
    if hasattr(sys.stdout, "buffer"):
        try:
            _console_stream = io.TextIOWrapper(
                sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True
            )
        except Exception:
            _console_stream = sys.stdout
_console_handler = logging.StreamHandler(_console_stream)
_console_handler.setFormatter(_log_formatter)

logging.basicConfig(level=logging.INFO, handlers=[_file_handler, _console_handler])
logger = logging.getLogger("create_qa_word")


GENERATION_PROMPT ="""
你是一名工业设备失效分析专家，熟悉核电、火电、石化、钢铁、船舶、矿山等场景中的通用设备失效分析，包括管道、阀门、水泵、风机、换热器、压力容器、紧固件、轴承、焊接接头、密封件等设备或部件。

你的任务是根据输入的技术报告文本，抽取高质量中文问答数据，用于训练“工业设备失效分析问答模型”。

请注意：你要生成的是“基于失效事实和检测证据进行工程判断”的问答数据，而不是阅读理解题、报告总结题、检测流程管理题或泛泛的方法论问题。

一、核心生成目标

你生成的每一条 QA 都必须围绕以下逻辑展开：

发现了什么事实情况 → 这些事实说明什么工程问题 → 可以支持什么失效结论 → 应提出什么工程建议

也就是说，问题和答案应重点训练模型完成以下能力：

1. 从故障现象中识别异常特征。
2. 从宏观检查、化学成分、金相检验、硬度试验、微观分析中提取关键证据。
3. 根据多项检测证据判断可能的失效机理。
4. 根据证据排除不符合的失效原因。
5. 在证据充分时给出失效结论。
6. 在证据不足时说明不能直接下结论，并提出进一步检测建议。
7. 根据失效机理提出工程整改、预防或运行维护建议。

二、必须生成的问题类型

请优先生成以下类型的问题：

1. 事实发现 → 机理判断型

问题应围绕“检测发现了什么异常现象，这些异常现象说明可能发生了什么失效机理”。

示例方向：
“某承压管道内壁发现点蚀坑，裂纹从内壁起裂并向外壁扩展，裂纹呈分叉、尖细和树枝状特征，同时材料成分和基体组织未见明显异常，应如何判断其主要失效机理？”

2. 证据组合 → 结论推导型

问题应要求模型综合多个检测结果，而不是只解释单个现象。

示例方向：
“当宏观检查显示裂纹起源于介质接触表面，金相检验显示裂纹与腐蚀坑相连，硬度和化学成分均符合要求时，应如何综合判断失效原因？”

3. 异常现象 → 风险判断型

问题应围绕某些检测事实说明设备存在什么风险。

示例方向：
“如果某设备部件表面存在较多腐蚀坑，且裂纹多从腐蚀坑处萌生，这对后续运行安全意味着什么？”

4. 检测证据 → 排除原因型

问题应体现如何根据证据排除材料错用、热处理异常、单纯过载、制造缺陷等原因。

示例方向：
“如果失效件化学成分符合标准、硬度无明显异常、基体组织正常，但裂纹均从内壁腐蚀坑处起裂，可以优先排除哪些失效原因？”

5. 结论 → 工程建议型

问题应围绕失效分析结论提出针对性建议。

示例方向：
“当多项检测结果支持腐蚀诱发开裂时，应从介质控制、表面防护、定期检测和运行工况管理等方面提出哪些工程建议？”

6. 证据不足 → 补充检测型

如果输入文本证据不足，不要强行生成失效机理结论，应生成进一步检测建议型问题。

示例方向：
“如果现有资料只显示部件存在裂纹和腐蚀痕迹，但缺少金相、硬度和微观断口分析，应如何安排后续检测以判断失效机理？”

三、禁止生成的问题类型

不得生成以下类型的问题：

1. 检测流程管理类问题

例如不得生成：
“多件相似失效样品分析中，如何处理以保证效率与代表性？”
“如何安排多个样品的检测顺序？”
“如何提高失效分析工作的效率？”

2. 报告写作类问题

例如不得生成：
“失效分析报告应如何组织章节？”
“报告中图表应该如何编号？”
“如何描述多个样品的相似性？”

3. 泛泛方法论问题

例如不得生成：
“失效分析通常包括哪些步骤？”
“宏观检查有什么作用？”
“金相检验为什么重要？”

除非输入证据不足，只能生成补充检测建议型 QA，否则不要生成这种泛泛的知识解释题。

4. 单纯阅读理解题

例如不得生成：
“该报告检测了哪些样品？”
“该文件中的表 5-1 说明了什么？”
“该报告第几章给出了结论？”
“某编号样品的检测结果是什么？”

5. 与失效结论无直接关系的问题

例如不得生成：
“多个样品形貌相似时如何提高分析代表性？”
“相似样品是否可以合并描述？”
“检测报告如何避免重复表述？”

四、问题生成范围

生成的问题必须围绕以下五类检测证据展开，并尽量与失效结论或工程建议关联。

1. 宏观检查

关注：
裂纹位置、起裂区域、断口形貌、腐蚀坑、磨损痕迹、变形、泄漏位置、表面颜色、沉积物、焊缝附近异常等。

问题应体现：
这些宏观事实提示了什么失效方向，例如腐蚀、疲劳、磨损、过载、制造缺陷或介质作用等。

2. 化学成分分析

关注：
材料成分是否符合标准、关键合金元素是否异常、是否存在材料错用、元素偏析、腐蚀介质残留等。

问题应体现：
化学成分结果如何支持或排除材料错用、材料不合格、耐蚀性不足等原因。

3. 金相检验

关注：
显微组织类型、夹杂物、晶粒状态、脱碳、过热、组织异常、裂纹扩展路径、晶间或穿晶特征、腐蚀坑与裂纹关系等。

问题应体现：
金相结果如何判断裂纹起源、扩展方式和可能失效机理。

4. 硬度试验

关注：
硬度是否异常、硬度分布是否均匀、是否存在局部硬化或软化、热处理状态是否异常、硬度与开裂、磨损、脆化之间的关系等。

问题应体现：
硬度结果如何支持或排除热处理异常、材料强度异常、加工硬化、局部脆化等原因。

5. 微观分析

关注：
SEM 断口形貌、韧窝、解理、疲劳条带、沿晶断裂、腐蚀产物、能谱分析、沉积物成分、氧化物、氯、硫等腐蚀性元素。

问题应体现：
微观形貌和微区成分如何支撑疲劳、脆断、腐蚀、磨损、氧化或介质作用等机理判断。

五、证据使用规则

1. 所有问答都必须基于输入文本中的证据。
2. 不得虚构输入中没有出现的材料牌号、温度、压力、介质、检测数据、设备编号、标准编号或实验结果。
3. 可以对具体设备、样品和编号进行抽象，但不能改变事实含义。
4. 如果输入中有具体编号，应抽象为“某工业设备部件”“某承压管道”“某换热器管束”“某阀门部件”“某焊接接头”等。
5. 如果输入中出现具体数值，可以在 answer 中概括其工程意义，但不要把 question 设计成单纯询问数值。
6. 如果证据不足，不得生成确定性失效结论，只能生成证据解释型、现象描述型或进一步检测建议型 QA。
7. thinking 中必须体现证据链，而不是简单重复答案。

六、抽象化要求

生成问题时必须进行工程抽象，避免针对原文细节发问。

请避免以下写法：

* “该报告中的 001VP 管失效原因是什么？”
* “图 5-1 显示了什么？”
* “表 5-1 的化学成分结果是什么？”
* “该文件第 3 章的结论是什么？”
* “本文中的样品 A 的硬度是多少？”

应改写为通用工程失效分析问题，例如：

* “某承压管道内壁存在点蚀坑，裂纹从点蚀坑处起裂并呈分叉扩展时，应如何判断其失效机理？”
* “当材料化学成分满足标准要求但仍发生裂纹泄漏时，应如何结合宏观和金相结果继续排查原因？”
* “如果硬度和组织均无明显异常，而裂纹起源于介质接触表面的腐蚀坑，应如何判断材料因素与环境因素在失效中的作用？”

七、输出数量要求

请根据输入文本的信息密度生成问答对。

* 如果输入内容较少，生成 3–5 条高质量 QA。
* 如果输入内容较完整，生成 5–10 条高质量 QA。
* 如果输入包含完整的宏观、化学成分、金相、硬度和微观分析内容，可以生成 10–15 条 QA。

质量优先，不要为了凑数量生成重复、空泛或证据不足的问题。

八、输出格式要求

你必须只输出合法 JSONL，不要输出 Markdown，不要输出 ```json 代码块，不要输出解释性文字。

每一行必须是一个独立 JSON 对象，不要输出 JSON 数组。

每个 JSON 对象格式如下：

{
"question": "面向通用工业设备失效分析场景的中文问题",
"thinking": [
"1. 发现的关键故障事实、检测结果或运行异常。",
"2. 该事实说明的工程含义，例如起裂位置、损伤类型、材料状态或环境作用。",
"3. 另一项关键检测证据及其工程含义。",
"4. 多项证据之间如何相互印证，并指向某类失效机理。",
"5. 根据现有证据可以排除或暂不能确认的其他原因。",
"6. 最终可支持的失效结论、风险判断、工程建议或进一步检测建议。"
],
"answer": "基于事实和证据得到的简洁中文结论或建议。"
}

九、字段要求

1. question

* 必须是中文。
* 必须围绕“事实情况 → 失效结论 → 工程建议”展开。
* 必须是通用工程设备失效分析问题。
* 不要出现具体文件名、图号、表号、样品编号、页码。
* 不要生成检测流程管理、报告写作、样品代表性或效率优化类问题。
* 尽量体现推理性，不要只问“是什么”。

2. thinking

* 必须是数组。
* 每条 thinking 必须包含证据和推理含义。
* 应体现“现象—证据—分析—排除—结论/建议”的过程。
* 不要编造输入中不存在的证据。
* 如果证据不足，应明确说明“现有证据不足以直接判断最终机理”。

3. answer

* 必须是中文。
* 应简洁、明确、工程化。
* 应回答“发现了什么事实、可以支持什么判断、应采取什么建议”。
* 不要写成泛泛的教材解释。
* 不要脱离输入内容泛泛而谈。

十、生成原则

请优先生成能够训练模型完成以下任务的 QA：

1. 根据宏观裂纹、腐蚀坑、断口、泄漏位置判断可能失效方向。
2. 根据化学成分分析判断材料是否符合要求，以及是否可以排除材料错用。
3. 根据金相组织和裂纹扩展路径判断裂纹起源和扩展方式。
4. 根据硬度结果判断材料状态、加工影响或热处理异常。
5. 根据微观断口和能谱结果判断疲劳、脆断、腐蚀、磨损、氧化或介质作用。
6. 综合多种检测结果判断失效机理。
7. 根据失效机理提出工程预防和整改建议。
8. 在证据不足时提出合理的补充检测方案。

十一、最终检查

输出前请逐条检查：

1. 这个问题是否围绕具体失效事实和检测证据？
2. 这个问题是否能训练模型判断失效原因、风险或工程建议？
3. 这个问题是否避免了报告写作、样品管理、流程效率、代表性分析等无关方向？
4. 这个问题是否避免了针对具体文件、图表、编号的阅读理解？
5. answer 是否是由输入证据支持的结论或建议？

只有同时满足以上要求的 QA 才能输出。
"""


def is_noise_line(text: str) -> bool:
    """
    判断一行文本是否属于噪声内容，例如：
    目录行、图题、表题、图片标注、页码、小标签等
    """

    text = text.strip()

    if not text:
        return True

    # 1. 过滤很短的图片标注词
    short_noise_words = {
        "宏观", "内壁", "外壁", "中部", "整体情况",
        "位置A", "位置B", "位置C", "位置D", "位置E", "位置F",
        "启裂处", "裂纹尖端", "中间位置",
        "001VP", "002VP",
        "区域1", "区域2", "区域3", "区域4"
    }

    if text in short_noise_words:
        return True

    # 2. 过滤图题，例如：图5-1 来样管外壁宏观检查
    if re.match(r"^图\s*\d+[-－]\d+", text):
        return True

    # 3. 过滤表题，例如：表5-1 化学成分分析结果
    if re.match(r"^表\s*\d+[-－]\d+", text):
        return True

    # 4. 过滤目录标题
    if text in {"目录", "目  录", "目 录"}:
        return True

    # 5. 过滤目录中的行，例如：5.1 宏观检查    7
    if re.match(r"^\d+(\.\d+)*\s+.+\s+\d+$", text):
        return True

    # 6. 过滤纯页码
    if re.match(r"^\d+$", text):
        return True

    # 7. 过滤类似：(a)、(b)、（a）、（b）
    if re.match(r"^[\(（][a-zA-Z0-9一二三四五六七八九十]+[\)）]", text):
        return True

    # 8. 过滤只有编号或位置编号的行，例如：位置1、位置10
    if re.match(r"^位置\d+$", text):
        return True

    # 9. 过滤过短且没有中文说明意义的行
    if len(text) <= 2:
        return True

    return False


def clean_text_lines(lines: list) -> list:
    """
    对提取出来的文本行进行清洗。
    """
    cleaned = []

    for line in lines:
        line = line.strip()

        if is_noise_line(line):
            continue

        # 合并多余空格
        line = re.sub(r"\s+", " ", line)

        cleaned.append(line)

    return cleaned


def extract_docx_content(docx_path: str, extract_tables: bool = True) -> str:
    """
    提取并清洗单个 docx 文件中的文本内容。

    参数:
        docx_path: docx 文件路径
        extract_tables: 是否提取表格内容，默认提取

    返回:
        str: 清洗后的完整文本内容
    """

    try:
        doc = Document(docx_path)
    except Exception as e:
        # 常见于：老的 .doc 被改名成 .docx、文件损坏或空文件。
        # python-docx 只支持 OOXML 格式的 .docx。
        raise ValueError(
            f"无法作为 .docx 打开（可能是老式 .doc 改名、损坏或空文件）：{e}"
        )
    content_parts = []

    # 1. 提取普通段落
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            content_parts.append(text)

    # 2. 提取表格内容
    if extract_tables:
        for table in doc.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text and not is_noise_line(cell_text):
                        row_text.append(cell_text)

                if row_text:
                    content_parts.append(" | ".join(row_text))

    # 3. 清洗噪声行
    cleaned_lines = clean_text_lines(content_parts)

    return "\n".join(cleaned_lines)


def is_section_title(line: str) -> bool:
    """
    判断清洗后的行是否像章节标题。
    """

    line = line.strip()

    if not line or len(line) > 50:
        return False

    title_patterns = [
        r"^第[一二三四五六七八九十百\d]+[章节部分篇]\s*",
        r"^[一二三四五六七八九十]+[、.．]\s*\S+",
        r"^\d+(\.\d+)*[、.．]?\s+\S+",
    ]

    return any(re.match(pattern, line) for pattern in title_patterns)


def split_text_by_sections(cleaned_text: str) -> list:
    """
    按章节组织 chunk：短章节合并，长章节单独作为一个 chunk。
    """

    lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]
    if not lines:
        return []

    sections = []
    current_section = []

    for line in lines:
        if is_section_title(line) and current_section:
            sections.append("\n".join(current_section))
            current_section = [line]
        else:
            current_section.append(line)

    if current_section:
        sections.append("\n".join(current_section))

    if len(sections) == 1:
        return split_text_by_max_chars(sections[0])

    chunks = []
    current_chunk = ""

    for section in sections:
        section_len = len(section)

        if section_len >= MIN_CHUNK_CHARS:
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""
            chunks.append(section)
            continue

        if not current_chunk:
            current_chunk = section
        elif len(current_chunk) + section_len + 1 <= MAX_CHUNK_CHARS:
            current_chunk = current_chunk + "\n" + section
        else:
            chunks.append(current_chunk)
            current_chunk = section

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def split_text_by_max_chars(text: str) -> list:
    """
    未识别出章节时，按 MAX_CHUNK_CHARS 顺序切分文本。
    """

    chunks = []
    current_chunk = ""

    for line in [line.strip() for line in text.splitlines() if line.strip()]:
        if len(line) > MAX_CHUNK_CHARS:
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""
            for start in range(0, len(line), MAX_CHUNK_CHARS):
                chunks.append(line[start:start + MAX_CHUNK_CHARS])
            continue

        if not current_chunk:
            current_chunk = line
        elif len(current_chunk) + len(line) + 1 <= MAX_CHUNK_CHARS:
            current_chunk = current_chunk + "\n" + line
        else:
            chunks.append(current_chunk)
            current_chunk = line

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def call_glm_generate(cleaned_text: str, max_retries: int = 3) -> str:
    """
    将清洗后的文本输入 GLM 模型，生成数据。带指数退避重试。

    参数:
        cleaned_text: 单个 chunk 清洗后的文本
        max_retries: 最大尝试次数（含首次）

    返回:
        str: GLM 模型生成结果（保证非空）

    异常:
        重试耗尽后抛出最后一次异常。
    """

    # 部分自定义/私有部署端点不支持 system role，遇到会返回非标准响应，
    # 故沿用单条 user 消息，将 prompt 与文本拼接（与原始可用版本保持一致）
    user_content = f"""
{GENERATION_PROMPT}

以下是清洗后的原始文本：

{cleaned_text}
"""
    messages = [
        {"role": "user", "content": user_content},
    ]

    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=GLM_MODEL,
                messages=messages,
                temperature=0,
                timeout=1800.0,  # 单次调用超时 30 分钟
            )

            content = resp.choices[0].message.content
            if content is None or not content.strip():
                raise ValueError("模型返回空内容")

            return content

        except Exception as e:
            last_err = e
            if attempt < max_retries:
                sleep_s = min(2 ** attempt, 30)
                logger.warning("模型调用失败（第 %d/%d 次）：%s，%d 秒后重试",
                               attempt, max_retries, e, sleep_s)
                time.sleep(sleep_s)

    raise last_err


def _strip_code_fence(text: str) -> str:
    """
    剥离模型输出中可能存在的 ```json / ``` 代码块围栏。
    """

    text = text.strip()

    if not text.startswith("```"):
        return text

    lines = text.splitlines()

    # 去掉首行的 ``` 或 ```json
    if lines and lines[0].lstrip().startswith("```"):
        lines = lines[1:]

    # 去掉末行的 ```
    if lines and lines[-1].strip().startswith("```"):
        lines = lines[:-1]

    return "\n".join(lines)


def parse_generated_records(generated_data: str, source_pdf: str, qa_type: str, chunk_id: int):
    """
    解析模型生成的 JSONL 文本为记录列表，逐行容错。

    返回:
        (records, bad_count): 解析成功的记录列表，以及解析失败的行数
    """

    records = []
    bad_count = 0

    cleaned = _strip_code_fence(generated_data)

    for line in cleaned.splitlines():
        line = line.strip()

        if not line or line.startswith("```"):
            continue

        try:
            item = json.loads(line)
        except json.JSONDecodeError:
            bad_count += 1
            logger.warning("[%s] chunk %d 跳过无法解析的行：%s",
                           source_pdf, chunk_id, line[:80])
            continue

        item["qa_type"] = qa_type
        item["source_pdf"] = source_pdf
        item["page_start"] = 0
        item["page_end"] = 0
        item["chunk_id"] = chunk_id

        records.append(item)

    return records, bad_count


def process_one_file(file_path: str, filename: str, qa_type: str, max_retries: int) -> dict:
    """
    处理单个 docx 文件：提取清洗 → 切 chunk → 逐 chunk 调用模型并解析。

    文件内所有 chunk 全部在内存收集，任一 chunk 重试耗尽仍失败会抛异常，
    从而保证“要么整文件成功、要么整文件失败”的原子性（不写半截数据）。

    返回:
        dict: {filename, records, bad, chunks, empty}
    """

    text = extract_docx_content(file_path, extract_tables=False)

    if not text.strip():
        return {"filename": filename, "records": [], "bad": 0, "chunks": 0, "empty": True}

    chunks = split_text_by_sections(text)
    logger.info("[%s] 清洗完成（%d 字），切分为 %d 个 chunk", filename, len(text), len(chunks))

    all_records = []
    total_bad = 0

    for chunk_id, chunk_text in enumerate(chunks, start=1):
        generated_data = call_glm_generate(chunk_text, max_retries=max_retries)
        records, bad = parse_generated_records(generated_data, filename, qa_type, chunk_id)
        all_records.extend(records)
        total_bad += bad

    return {
        "filename": filename,
        "records": all_records,
        "bad": total_bad,
        "chunks": len(chunks),
        "empty": False,
    }


def append_records(output_path: str, records: list):
    """
    将记录列表追加写入 jsonl 文件（由主线程串行调用，无需加锁）。
    """

    with open(output_path, "a", encoding="utf-8") as f:
        for item in records:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")


def load_done_set(done_path: str) -> set:
    """
    读取已完整处理文件清单。
    """

    if not os.path.exists(done_path):
        return set()

    with open(done_path, "r", encoding="utf-8") as f:
        return {line.strip() for line in f if line.strip()}


def mark_done(done_path: str, filename: str):
    """
    将一个已完整处理的文件追加到清单。
    """

    with open(done_path, "a", encoding="utf-8") as f:
        f.write(filename + "\n")


def collect_source_pdfs_from_jsonl(output_path: str) -> set:
    """
    从已有 jsonl 结果中收集出现过的 source_pdf，用于兼容旧数据的断点迁移。
    """

    result = set()

    if not os.path.exists(output_path):
        return result

    with open(output_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                item = json.loads(line)
            except json.JSONDecodeError:
                continue
            src = item.get("source_pdf")
            if src:
                result.add(src)

    return result


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input_folder", default="papers_failure_analysis2")
    parser.add_argument("--output_file", default="qa_word.jsonl")
    parser.add_argument("--error_file", default="error_log.txt")
    parser.add_argument("--qa_type", default="")
    parser.add_argument("--write_mode", choices=["resume", "rewrite"], default="resume")
    parser.add_argument("--max_workers", type=int, default=4, help="文件间并发数")
    parser.add_argument("--max_retries", type=int, default=3, help="单次模型调用最大尝试次数")
    args = parser.parse_args()

    # 启动诊断：定位工作目录与输入目录，便于排查“路径/挂载”类卡顿或读不到文件
    logger.info("脚本启动 [build=utf8-fix-2]，stdout 编码：%s，工作目录：%s",
                getattr(sys.stdout, "encoding", "?"), os.getcwd())
    logger.info("输入目录参数：%s -> 绝对路径：%s（存在：%s）",
                args.input_folder, os.path.abspath(args.input_folder),
                os.path.isdir(args.input_folder))

    # 配置校验：尽早发现 .env 未配置的问题，而不是等第一次 API 调用才失败
    if not GLM_API_KEY or not GLM_BASE_URL:
        logger.error("API_KEY / API_BASE_URL 未配置，请检查 .env 文件")
        return

    input_folder = args.input_folder
    output_file = args.output_file
    error_file = args.error_file
    done_file = output_file + ".done"  # 已完整处理文件清单

    if not os.path.isdir(input_folder):
        logger.error("输入目录不存在：%s", input_folder)
        return

    if args.write_mode == "rewrite":
        confirm = input(f"确认要覆盖 {output_file} 吗？输入 'yes' 确认：")
        if confirm.lower() != "yes":
            logger.info("操作已取消")
            return
        open(output_file, "w", encoding="utf-8").close()
        open(done_file, "w", encoding="utf-8").close()  # 同步清空进度清单

    # 固定顺序，保证多次运行文件列表一致
    filenames = sorted(
        filename for filename in os.listdir(input_folder)
        if filename.lower().endswith(".docx") and not filename.startswith("~$")
    )

    done_set = set()
    if args.write_mode == "resume":
        done_set = load_done_set(done_file)

        # 兼容旧数据：清单不存在但已有 jsonl 结果时，用其中的 source_pdf 初始化进度
        if not done_set and os.path.exists(output_file):
            migrated = collect_source_pdfs_from_jsonl(output_file)
            if migrated:
                for name in sorted(migrated):
                    mark_done(done_file, name)
                done_set = migrated
                logger.info("检测到旧 jsonl 结果，已迁移 %d 个文件到进度清单", len(migrated))

    todo = [f for f in filenames if f not in done_set]

    if not todo:
        logger.info("没有需要处理的文件（共 %d 个，已完成 %d 个）", len(filenames), len(done_set))
        return

    logger.info("待处理 %d 个文件（已完成 %d 个），并发数 %d",
                len(todo), len(done_set), args.max_workers)

    grand_total = 0

    # 文件间并发；每个 worker 返回整文件结果，由主线程串行写入并登记，保证写入原子性与断点可靠
    with ThreadPoolExecutor(max_workers=args.max_workers) as executor:
        future_to_name = {
            executor.submit(
                process_one_file,
                os.path.join(input_folder, filename),
                filename,
                args.qa_type,
                args.max_retries,
            ): filename
            for filename in todo
        }

        for future in as_completed(future_to_name):
            filename = future_to_name[future]
            try:
                result = future.result()

                if result["empty"]:
                    logger.warning("[%s] 清洗后文本为空，跳过", filename)
                    mark_done(done_file, filename)  # 空文件也登记，避免重复处理
                    continue

                append_records(output_file, result["records"])
                mark_done(done_file, filename)
                grand_total += len(result["records"])

                bad_note = f"，丢弃坏行 {result['bad']} 条" if result["bad"] else ""
                logger.info("[%s] 完成，写入 %d 条数据%s", filename, len(result["records"]), bad_note)

            except Exception as e:
                logger.error("[%s] 处理失败：%s", filename, e)
                with open(error_file, "a", encoding="utf-8") as f:
                    f.write(f"{filename}\t{str(e)}\n")

    logger.info("全部结束，本次共写入 %d 条数据", grand_total)


if __name__ == "__main__":
    main()
